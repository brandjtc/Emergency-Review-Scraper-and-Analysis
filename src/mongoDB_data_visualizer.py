import matplotlib.pyplot as plt
import numpy as np
import os
import pymongo
from pymongo import MongoClient
import pandas as pd
import datetime as dt
from tzlocal import get_localzone
import pandas as pd
from docx import Document as doc
import random
from dotenv import load_dotenv, find_dotenv
_ = load_dotenv(find_dotenv())

#Establishes connection to mongo DB
mongoConnectStr=os.getenv("MONGO_KEY")
client = MongoClient(mongoConnectStr)
#Selects EmergencyReviewDB as the database in Mongo to populate
dbString=os.environ.get("DATABASE")
proj_db=client[dbString]

#Connects to tables in aforementioned DB
colString=os.environ.get("COLLECTION")
review_collection = proj_db[colString]

#Variable instantiation for time functions
start= dt.datetime.now(tz=get_localzone())

def timeStart():
    print("-"*40)
    fmt = "%m/%d/%y - %T %p"  
    start= dt.datetime.now(tz=get_localzone())
    print("Began at:"+start.strftime(fmt))

def timeEnd():
    print('-'*40)
    fmtend = "%m/%d/%y - %T %p"  
    end= dt.datetime.now(tz=get_localzone())
    print("Ended at: "+end.strftime(fmtend))
    print(f"Time elapsed: {end - start}")
    print("-"*40)
    print("\n")

#Variable declaration for data gathering, used in report doc generation later
playStoreReviewCount=0
playStorePercent=0
appStoreReviewCount=0
appStorePercent=0
meanLength=0
standardDeviation=0
shortLenList=[]
shortCount=0
shortReviewPercent=0
middleCount=0
middleReviewPercent=0
largeCount=0
largeReviewPercent=0
sampledShortRev=[]
print("Converting input excel info to lists and calculating total apps and generating figure")
timeStart()

#Gathering Data from input excel
ios_excel=pd.read_excel(os.getenv("EXCELFILEPATH"),usecols='L,M')
appStoreExcelData = list(ios_excel['Link to App Store'])
playStoreExcelData= list(ios_excel['Link to Google Play Store'])

#Variable declaration for graphic generation
playstoreAndIos=0
exclusivePlaystore=0
exclusiveAppstore=0
for i in range(0,len(appStoreExcelData)):
    if (str(appStoreExcelData[i])!="nan")&(str(playStoreExcelData[i])!="nan"):
        playstoreAndIos+=1
    elif (str(appStoreExcelData[i])!="nan")&(str(playStoreExcelData[i])=="nan"):
        exclusiveAppstore+=1
    elif (str(playStoreExcelData[i])!="nan")&(str(appStoreExcelData[i])=="nan"):
        exclusivePlaystore+=1


barXinp=np.array(['Exclusive to\nGoogle PlayStore','Exclusive to\nIOS App Store','IOS and\n Google Play'])
barYinp=np.array([exclusivePlaystore,exclusiveAppstore,playstoreAndIos])
fig, ax = plt.subplots()
bars = ax.bar(barXinp, barYinp)

ax.bar_label(bars)

plt.title("App distribution\n among App Stores")
plt.ylabel("Amount on specified platform")
plt.savefig('./Generated Files/storeCount.png', 
           transparent=False,  
           facecolor='white', 
           bbox_inches="tight")



timeEnd()

print("Gathering review percentage counts")
timeStart()
shortReviewList=list()
data = review_collection.find({},{"_id":0,"content":1,"platform":1})
for item in data:
    if(item["content"]!=None):
        if(len(item["content"])>25):
            if (item["platform"]=="Google Play Store"):
                playStoreReviewCount+=1
            else:
                appStoreReviewCount+=1
            if(len(item["content"])<200):
                middleCount+=1
            else:
                largeCount+=1
        else:
            shortCount+=1
            shortLenList.append(len(item["content"]))
            shortReviewList.append(item["content"])
    else:
        shortCount+=1

preRoundRevTotal=middleCount+shortCount+largeCount

middleReviewPercent=round((middleCount/preRoundRevTotal)*100,2)
largeReviewPercent=round((largeCount/preRoundRevTotal)*100,2)

reviewTotal=playStoreReviewCount+appStoreReviewCount
playStorePercent=round((playStoreReviewCount/reviewTotal)*100,2)
appStorePercent=round((appStoreReviewCount/reviewTotal)*100,2)

timeEnd()


print("Counting length of all reviews in the database and generating figure")
timeStart()

#Retrieving all reviews from mongoDB
data = review_collection.find({},{"_id":1,"content":1})
histogramInput = []
for item in data:
    try:
        histogramInput.append(len(item["content"]))
    except TypeError:
            pass
    

#Gathers the percentage of short reviews. These will all be excluded from final dataset for the rest of these calculations.
shortReviewPercent=round(shortCount/preRoundRevTotal*100,2)


client.close()
# Creation of histogram containing review lengths as dataset
fig, ax = plt.subplots()
values, bins, bars = plt.hist(np.log10(histogramInput),linewidth=0.5, edgecolor="white")
plt.bar_label(bars)
plt.margins(x=0.1,y=0.1)

plt.title("Review Length Frequency")
plt.ylabel("Amount")
plt.xlabel("Log (Base 10) of Review Length in Characters")
plt.savefig('./Generated Files/reviewLength_Histogram.png', 
           transparent=False,  
           facecolor='white', 
           bbox_inches="tight")
timeEnd()

print("Building report doc")
timeStart()
#Histogram input already has the short lengths removed meaning it can be
#used for calculating standard deviation and mean
standardDeviation=round(np.std(histogramInput),2)
meanLength=round(np.mean(histogramInput),2)
sampledShortRev=random.sample(shortReviewList,3)

reportDoc = doc()

reportDoc.add_heading('Scraper Report', 0)
reportDoc.add_heading('Summary',level=1)
summaryString=f'''The length of the reviews varied (M={meanLength}, SD={standardDeviation}). Approximately {shortReviewPercent}% consisted of just a few words (<25 characters), Approximately {middleReviewPercent}% contained at most one sentence (<200 characters), while {largeReviewPercent}% contained multiple sentences (>201 characters). Since reviews shorter than a few words (<25 characters) likely contain trivial comments such as "{sampledShortRev[0]}", "{sampledShortRev[1]}", and "{sampledShortRev[2]}." Their content is unlikely to have the depth or breadth needed to significantly invoke any of the dimensions or sub-dimensions presented in our theoretical model. Therefore, we excluded them from the dataset. Thus, after removing {shortCount} reviews, our final dataset consisted of {reviewTotal} reviews, with approximately ({playStorePercent}% being from the Google Play Store and {appStorePercent}% from the Apple AppStore).'''
reportDoc.add_paragraph(summaryString)
reportDoc.add_heading('Figure 1', level=1)
reportDoc.add_picture('./Generated Files/storeCount.png')
reportDoc.add_heading('Figure 2', level=1)
reportDoc.add_picture('./Generated Files/reviewLength_Histogram.png')
reportDoc.save('./Generated Files/report.docx')
timeEnd()
client.close()