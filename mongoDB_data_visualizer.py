import matplotlib.pyplot as plt
import numpy as np
import os
import pymongo
from pymongo import MongoClient
import pandas as pd
import datetime as dt
from tzlocal import get_localzone
import pandas as pd
from docx import Document as doc

from dotenv import load_dotenv, find_dotenv
_ = load_dotenv(find_dotenv())

#Establishes connection to mongo DB
client = MongoClient(os.getenv("MONGO_KEY"))
#Selects EmergencyReviewDB as the database in Mongo to populate
dbString=os.environ.get("DATABASE")
proj_db=client[str(dbString)]

#Connects to tables in aforementioned DB
colString=os.environ.get("COLLECTION")
review_collection = proj_db[str(colString)]

#Variable instantiation for time functions
start= dt.datetime.now(tz=get_localzone())

def timeStart():
    print("-"*40)
    fmt = "%m/%d/%y - %T %p"  
    start= dt.datetime.now(tz=get_localzone())
    print("Began at:"+start.strftime(fmt))

def timeEnd():
    print('-'*40)
    fmtend = "%m/%d/%y - %T %p"  
    end= dt.datetime.now(tz=get_localzone())
    print("Ended at: "+end.strftime(fmtend))
    print(f"Time elapsed: {end - start}")
    print("-"*40)
    print("\n")

print("Counting app reviews.")
timeStart()


playstoreReviewCount=proj_db.review_collection.count_documents({"platform":"Google Play Store"})
iosReviewCount= proj_db.review_collection.count_documents({"platform":"Apple App Store"})

timeEnd()

print("Converting input excel info to lists and calculating total apps and generating figure")
timeStart()

ios_excel=pd.read_excel(os.getenv("EXCELFILEPATH"),usecols='L,M')
appStoreExcelData = list(ios_excel['Link to App Store'])
playStoreExcelData= list(ios_excel['Link to Google Play Store'])

playstoreAndIos=0
exclusivePlaystore=0
exclusiveAppstore=0
for i in range(0,len(appStoreExcelData)):
    if (str(appStoreExcelData[i])!="nan")&(str(playStoreExcelData[i])!="nan"):
        playstoreAndIos+=1
    elif (str(appStoreExcelData[i])!="nan")&(str(playStoreExcelData[i])=="nan"):
        exclusiveAppstore+=1
    elif (str(playStoreExcelData[i])!="nan")&(str(appStoreExcelData[i])=="nan"):
        exclusivePlaystore+=1


barXinp=np.array(['Exclusive to\nGoogle PlayStore','Exclusive to\nIOS App Store','IOS and\n Google Play'])
barYinp=np.array([exclusivePlaystore,exclusiveAppstore,playstoreAndIos])
fig, ax = plt.subplots()
bars = ax.bar(barXinp, barYinp)

ax.bar_label(bars)

plt.title("App distribution\n among App Stores")
plt.ylabel("Amount on specified platform")
plt.savefig('./Generated Files/storeCount.png', 
           transparent=False,  
           facecolor='white', 
           bbox_inches="tight")



timeEnd()

print("Counting length of all reviews in the database and generating figure")
timeStart()

#Retrieving all reviews from mongoDB
data = review_collection.find({},{"_id":1,"content":1})

histogramInput = []
for item in data:
    if 'content' in item:
        try:
            histogramInput.append(len(item["content"]))
            if len(item["content"])>4000:
                print(item["_id"])
        except TypeError:
                pass
        


client.close()
# Creation of histogram containing review lengths as dataset
fig, ax = plt.subplots()
values, bins, bars = plt.hist(np.log10(histogramInput),linewidth=0.5, edgecolor="white")
plt.bar_label(bars)
plt.margins(x=0.1,y=0.1)

plt.title("Review Length Frequency")
plt.ylabel("Amount")
plt.xlabel("Log (Base 10) of Review Length in Characters")
plt.savefig('./Generated Files/reviewLength_Histogram.png', 
           transparent=False,  
           facecolor='white', 
           bbox_inches="tight")
timeEnd()

print("Building report doc")
timeStart()

reportDoc = doc()

reportDoc.add_heading('Scraper Report', 0)

reportDoc.add_heading('Figure 1', level=1)
reportDoc.add_picture('./Generated Files/storeCount.png')
reportDoc.add_heading('Figure 2', level=1)
reportDoc.add_picture('./Generated Files/reviewLength_Histogram.png')
reportDoc.save('./Generated Files/report.docx')
timeEnd()
client.close()